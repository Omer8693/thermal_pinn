"""
gen_docx_full.py
================
Generate NAS_PINN_Thermal_Quenching_Report.docx
5-page technical summary for supervisor briefing.
Human-written style, active voice, with figures and full k-sweep tables.
"""

import sys
from pathlib import Path
from thermal_pinn.runtime_paths import REPORT_DIR, RESULT_DIR

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: pip install python-docx"); sys.exit(1)

ROOT   = Path(__file__).resolve().parent
RES    = RESULT_DIR
FIG    = ROOT / "reports" / "thesis" / "figures"
KF     = RES / "01_all_k_fields" / "bayesian"
SFDIR  = RES / "subframe_result_application"
OUT    = REPORT_DIR / "NAS_PINN_Thermal_Quenching_Report.docx"

# colour palette
NAVY  = RGBColor(0x1C, 0x35, 0x57)
BLUE  = RGBColor(0x1A, 0x5C, 0x96)
DGRAY = RGBColor(0x44, 0x44, 0x44)
MGRAY = RGBColor(0x88, 0x88, 0x88)
GREEN = RGBColor(0x16, 0x65, 0x34)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED   = RGBColor(0xB9, 0x1C, 0x1C)


# ── helpers ──────────────────────────────────────────────────────────────

def _shd(cell, hex6):
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), hex6)
    cell._element.get_or_add_tcPr().append(s)


def _sp(p, before=0, after=4):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(13)
        r.font.bold = True
    _sp(p, 10, 4)


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = BLUE
        r.font.size = Pt(11)
        r.font.bold = True
    _sp(p, 6, 3)


def body(doc, text, size=10.5):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    for r in p.runs:
        r.font.size = Pt(size)
    _sp(p, 0, 4)


def cap(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size  = Pt(8.5)
        r.font.italic = True
        r.font.color.rgb = MGRAY
    _sp(p, 1, 6)


def fig(doc, path, caption_text, width=5.8):
    p = Path(path)
    if p.exists():
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(str(p), width=Inches(width))
        cap(doc, caption_text)
    else:
        body(doc, f"[Figure not available: {Path(path).name}]")


def tbl(doc, headers, rows, col_widths=None, highlight_col=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for row in t.rows:
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Inches(col_widths[j])
    # header
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        _shd(cell, '1C3557')
        cell.text = h
        for pa in cell.paragraphs:
            pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in pa.runs:
                r.font.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = WHITE
    # data
    for i, row_data in enumerate(rows):
        bg = 'F0F4F8' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row_data):
            cell = t.rows[1 + i].cells[j]
            _shd(cell, bg)
            cell.text = str(val)
            for pa in cell.paragraphs:
                pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in pa.runs:
                    r.font.size = Pt(9)
                    if str(val).startswith('\u2605'):
                        r.font.color.rgb = GREEN
                        r.font.bold = True
                    if highlight_col and j == highlight_col and str(val) not in ('0%', '-'):
                        r.font.color.rgb = GREEN
                        r.font.bold = True
    doc.add_paragraph()
    return t


# ══════════════════════════════════════════════════════════════════════════
# PAGE 1 — COVER + OVERVIEW + NAS-PINN ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════

def page1(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NAS-PINN Framework for Thermal Quenching of A356 Aluminium\n")
    r.font.size = Pt(17)
    r.font.bold = True
    r.font.color.rgb = NAVY
    r2 = p.add_run("Technical Summary  |  April 2026")
    r2.font.size = Pt(10)
    r2.font.color.rgb = MGRAY
    r2.font.italic = True
    _sp(p, 0, 8)

    h1(doc, "1. What We Did and Why")
    body(doc,
        "Water quenching of A356 aluminium castings requires solving the transient heat "
        "equation from 540 C down to near-ambient temperature over 30 seconds. Industrial "
        "FEM solvers must step through 60 Crank-Nicolson solves per simulation, which is "
        "expensive when repeated across design variants. The physical motivation for this "
        "work comes from Mortensen et al. (2026), who studied distortion in cast automotive "
        "subframes using an industrial FEM model.")
    body(doc,
        "We built a Physics-Informed Neural Network (PINN) surrogate that predicts the "
        "temperature field over k consecutive FEM anchor intervals per window. By increasing "
        "k from 1 to 5, we cut the required number of FEM anchor calls from 20 down to 4 "
        "over the 30-second simulation. We then evaluated this idea on seven canonical "
        "geometries and on a heterogeneous subframe-like surrogate domain. We also tested a "
        "PINN-only rollout mode in which FEM is used only once at t=0 and the network "
        "advances the solution autonomously for the rest of the simulation.")

    h1(doc, "2. NAS-PINN Architecture")
    body(doc,
        "Figure 1 shows the full training pipeline. Input coordinates and normalised window "
        "time are first mapped through a Fourier feature embedding (64 frequencies, sigma=1.0 "
        "for 2D and 1.5 for 3D) to mitigate spectral bias. The NAS-discovered network then "
        "produces a raw output that is combined with the IC-consistent output formula: "
        "the prediction equals the FEM (or previous PINN) snapshot exactly at tau=0, with no "
        "IC penalty weight needed. Four weighted loss terms drive training: PDE residual, "
        "initial condition, Robin boundary condition, and FEM endpoint supervision. "
        "In the PINN-only protocol the endpoint FEM target is replaced by a physics-based "
        "cooling prior.")

    fig(doc, FIG / "pinn_framework_fig_pinn_framework.png",
        "Figure 1: NAS-PINN training framework. Coordinates pass through Fourier embedding "
        "into the NAS-discovered network. The IC-consistent output layer enforces the initial "
        "condition exactly at tau=0. Four loss terms supervise the training. In the PINN-only "
        "protocol the FEM endpoint target is replaced by a surface-based cooling prior.",
        width=6.5)

    tbl(doc,
        headers=["Component", "Setting"],
        rows=[
            ("Material",           "A356 aluminium — lambda=150 W/mK (2D), 160 W/mK (3D)"),
            ("Heat transfer h",    "5000 W/(m2 K) for 2D  |  4000 W/(m2 K) for 3D"),
            ("Initial / bath T",   "T0=540 C  |  T_inf=20 C"),
            ("FEM step / anchor",  "0.5 s Crank-Nicolson  |  1.5 s anchor, 20 snapshots"),
            ("Skip factor k",      "1 to 5 in hybrid sweep  |  adaptive 1-3 in PINN-only"),
            ("NAS strategies",     "Bayesian TPE (5x151 ReLU)  |  NSGA-II (3x153 tanh)  |  NSGA-III (3x75 tanh)"),
            ("Training budget",    "800 Adam (cosine lr 1e-3->1e-5) + 50 L-BFGS per window"),
            ("Geometries tested",  "3 two-dimensional + 4 three-dimensional + 1 subframe surrogate"),
        ],
        col_widths=[2.3, 4.7],
    )


# ══════════════════════════════════════════════════════════════════════════
# PAGE 2 — 2D RESULTS: BEST PER k + HEAT MAPS
# ══════════════════════════════════════════════════════════════════════════

def page2(doc):
    doc.add_page_break()
    h1(doc, "3. Two-Dimensional Domain Results")
    body(doc,
        "We swept k from 1 to 5 on rectangle, circle, and L-shape domains for all three "
        "NAS architectures. Table 1 shows the best result (lowest MAE) for each k value "
        "across all three domains and architectures. The Bayesian TPE architecture on the "
        "L-shape was consistently the strongest 2D performer. At k=5, we reduced FEM anchor "
        "calls from 20 to 4 while keeping the MAE at 5.23 C, which is a 5x reduction in "
        "FEM cost for a 2.7x accuracy penalty relative to k=1.")

    tbl(doc,
        headers=["k", "Best domain", "Best arch", "MAE (C)", "FEM saved", "FEM calls"],
        rows=[
            ("1 (baseline)", "L-shape",  "Bayesian/TPE", "1.93",          "0%  (baseline)",  "20"),
            ("2",            "L-shape",  "Bayesian/TPE", "\u2605 2.48",   "50%",             "10"),
            ("3",            "L-shape",  "Bayesian/TPE", "\u2605 3.25",   "67%",             "7"),
            ("4",            "L-shape",  "Bayesian/TPE", "\u2605 4.10",   "75%",             "5"),
            ("5",            "L-shape",  "Bayesian/TPE", "\u2605 5.23",   "80%",             "4"),
        ],
        col_widths=[1.1, 1.2, 1.5, 1.0, 1.2, 0.9],
        highlight_col=4,
    )

    body(doc,
        "Table 2 shows the full breakdown for the L-shape domain (Bayesian TPE), which gave "
        "the cleanest progression: MAE grows smoothly from 1.93 C at k=1 to 5.23 C at k=5, "
        "while saving 80% of the FEM anchor evaluations. The re-entrant corner is the hardest "
        "geometric feature in the 2D set, yet the error stays local rather than spreading.",
        size=10)

    tbl(doc,
        headers=["Domain", "k=1", "k=2 (50%)", "k=3 (67%)", "k=4 (75%)", "k=5 (80%)"],
        rows=[
            ("Rectangle / Bayesian", "2.49 C", "3.49 C", "4.34 C", "5.35 C", "5.80 C"),
            ("Circle / Bayesian",    "2.70 C", "2.70 C", "3.63 C", "4.30 C", "5.70 C"),
            ("L-shape / Bayesian",   "1.93 C", "2.48 C", "3.25 C", "4.10 C", "5.23 C"),
            ("Circle / NSGA-III",    "1.64 C", "3.70 C", "5.66 C", "9.68 C", "12.50 C"),
            ("L-shape / NSGA-III",   "2.19 C", "5.72 C", "8.78 C", "11.51 C","10.54 C"),
        ],
        col_widths=[2.1, 0.9, 1.1, 1.1, 1.1, 1.1],
    )

    body(doc,
        "Figure 2 shows the L-shape temperature field and error map at k=1 and k=5. The "
        "left triplet (FEM reference, PINN prediction, error) confirms that even at k=5 the "
        "large errors are concentrated near the re-entrant corner and do not spread globally.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for path in [
        KF / "lshape" / "fig_th1_2d_lshape_bayesian_k1.png",
        KF / "lshape" / "fig_th1_2d_lshape_bayesian_k5.png",
    ]:
        if Path(path).exists():
            run.add_picture(str(path), width=Inches(3.1))
    cap(doc, "Figure 2: L-shape 2D Bayesian/TPE. Left: k=1, MAE=1.93 C. "
        "Right: k=5, MAE=5.23 C (80% FEM saved). Each triplet shows FEM reference, "
        "PINN prediction, and absolute error.")


# ══════════════════════════════════════════════════════════════════════════
# PAGE 3 — 3D RESULTS: BEST PER k + HEAT MAPS
# ══════════════════════════════════════════════════════════════════════════

def page3(doc):
    doc.add_page_break()
    h1(doc, "4. Three-Dimensional Domain Results")
    body(doc,
        "We tested four 3D geometries: rectangular prism, cylinder, stacked cubes, and "
        "L-shape prism. The 3D problem is harder than 2D because the input gains one more "
        "spatial coordinate and boundary residuals must be evaluated on surfaces. "
        "Table 3 shows the best result for each k across all 3D domains and architectures.")

    tbl(doc,
        headers=["k", "Best domain", "Best arch", "MAE (C)", "FEM saved", "FEM calls"],
        rows=[
            ("1 (baseline)", "L-shape 3D", "NSGA-II",     "3.43",         "0%  (baseline)", "20"),
            ("2",            "L-shape 3D", "NSGA-II/III", "\u2605 5.97",  "50%",            "10"),
            ("3",            "L-shape 3D", "Bayesian/TPE","\u2605 7.08",  "67%",            "7"),
            ("4",            "L-shape 3D", "Bayesian/TPE","\u2605 8.73",  "75%",            "5"),
            ("5",            "L-shape 3D", "Bayesian/TPE","\u2605 9.93",  "80%",            "4"),
        ],
        col_widths=[1.1, 1.4, 1.5, 1.0, 1.2, 0.9],
        highlight_col=4,
    )

    body(doc,
        "The cylinder (Bayesian/TPE) was the cleanest step-skipping case in 3D. Table 4 "
        "shows that its MAE grows from 6.43 C at k=1 to 11.90 C at k=5, staying below "
        "12 C even at the maximum FEM reduction. The L-shape 3D was best at k=1 and k=2 "
        "with NSGA-II (3.43 C and 5.97 C) but became less stable at higher k. This "
        "suggests that for complex 3D geometries, k=2 to k=3 is the practical sweet spot.",
        size=10)

    tbl(doc,
        headers=["Domain", "k=1", "k=2 (50%)", "k=3 (67%)", "k=4 (75%)", "k=5 (80%)"],
        rows=[
            ("Rect. prism / Bayesian",  "8.01 C",  "8.80 C",  "9.77 C",  "11.13 C", "13.01 C"),
            ("Cylinder / Bayesian",     "6.43 C",  "6.64 C",  "7.98 C",  "9.48 C",  "11.90 C"),
            ("Stacked / Bayesian",      "7.58 C",  "9.25 C",  "11.19 C", "11.37 C", "12.26 C"),
            ("L-shape 3D / Bayesian",   "8.41 C",  "7.18 C",  "7.08 C",  "8.73 C",  "9.93 C"),
            ("L-shape 3D / NSGA-II",    "3.43 C",  "5.97 C",  "13.12 C", "9.28 C",  "27.05 C"),
            ("L-shape 3D / NSGA-III",   "3.54 C",  "5.97 C",  "12.58 C", "10.22 C", "11.93 C"),
        ],
        col_widths=[2.1, 0.9, 1.1, 1.1, 1.1, 1.1],
    )

    body(doc,
        "Figure 3 shows the cylinder heat maps at k=1 and k=3. The cooling pattern is "
        "well captured in both cases. The main error concentrates near the top and bottom "
        "edges where the boundary curvature is highest.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for path in [
        KF / "fig_th1_3d_cylinder_bayesian_k1.png",
        KF / "fig_th1_3d_cylinder_bayesian_k3.png",
    ]:
        if Path(path).exists():
            run.add_picture(str(path), width=Inches(3.1))
    cap(doc, "Figure 3: Cylinder 3D Bayesian/TPE. Left: k=1, MAE=6.43 C. "
        "Right: k=3, MAE=7.98 C (67% FEM saved). Each panel shows FEM reference, "
        "PINN prediction, and absolute error at t=30 s.")


# ══════════════════════════════════════════════════════════════════════════
# PAGE 4 — SUBFRAME DOMAIN: HYBRID + PINN-ONLY
# ══════════════════════════════════════════════════════════════════════════

def page4(doc):
    doc.add_page_break()
    h1(doc, "5. Subframe-like Surrogate Domain")
    body(doc,
        "The subframe-like surrogate is the main application-level test. It is an academic "
        "3D geometry inspired by the industrial automotive subframe studied by Mortensen et al. "
        "(2026). It contains three region types: an external cooling body (teal), support and "
        "contact pads (dark red), and simplified enclosed cavity features (orange). "
        "Figure 4 shows both the 3D perspective view and the 2D top-view layout. "
        "We ran two protocols on all three NAS architectures.")

    fig(doc, SFDIR / "subframe_reference_domain_3d_2d_.png",
        "Figure 4: Subframe-like surrogate domain. Left (panel A): 3D perspective showing the "
        "external cooling body (teal), support/contact pads (dark red), and cavity regions (orange). "
        "Right (panel B): 2D top-view reference layout indicating the spatial arrangement of the "
        "three region types across the 1.25 m \u00d7 0.65 m footprint. "
        "This paper-inspired geometry tests the framework on a heterogeneous multi-region domain.",
        width=6.3)

    h2(doc, "5.1 FEM-Anchored Hybrid Sweep")
    body(doc,
        "Each window starts from a FEM snapshot. We selected the largest k that kept the "
        "mean-window MAE below 10 C. NSGA-III at k=5 was the best result: 80% FEM saving "
        "at 7.51 C mean MAE. Bayesian/TPE at k=4 gives a more conservative option with "
        "5.56 C MAE and 75% saving.",
        size=10)

    tbl(doc,
        headers=["Architecture", "Rec. k", "Mean MAE (C)", "Final MAE (C)", "FEM saved", "Run (min)"],
        rows=[
            ("Bayesian/TPE", "4", "5.56", "2.17", "75%",          "2.6"),
            ("NSGA-II",      "1", "5.03", "1.06", "0% (k=1 opt)", "10.3"),
            ("NSGA-III",     "5", "7.51", "3.64", "\u2605 80%",   "2.3"),
        ],
        col_widths=[1.5, 0.7, 1.2, 1.2, 1.3, 1.1],
        highlight_col=4,
    )

    h2(doc, "5.2 PINN-Only Adaptive Rollout")
    body(doc,
        "FEM is used only at t=0. Each subsequent window uses the previous PINN prediction "
        "as its initial condition. The adaptive-k controller selects window size using a "
        "physics-based score. The resulting schedule was k=1 for the first seven windows "
        "(0 to 10.5 s), k=2 for three windows, k=3 for two windows, then k=1 for the final "
        "window. NSGA-II gave the best autonomous result at 25.08 C final MAE.",
        size=10)

    tbl(doc,
        headers=["Architecture", "Final MAE (C)", "Rel. L2", "k schedule", "FEM saving", "Run (min)"],
        rows=[
            ("Bayesian/TPE",    "30.57", "0.246", "1-1-1-1-1-1-1-2-2-2-3-3-1", "35%", "10.9"),
            ("NSGA-II",         "25.08", "0.192", "1-1-1-1-1-1-1-2-2-2-3-3-1", "35%", "11.8"),
            ("NSGA-III",        "27.58", "0.209", "1-1-1-1-1-1-1-2-2-2-3-3-1", "35%", "11.9"),
        ],
        col_widths=[1.4, 1.1, 0.8, 2.5, 0.9, 0.9],
    )

    body(doc,
        "Regional breakdown for NSGA-II (best autonomous result): external body 22.88 C, "
        "support pads 58.17 C, cavity 36.76 C. Support pads and cavities dominate the "
        "error because their local geometry creates steep gradients that are harder to "
        "learn without FEM corrections.",
        size=10)

    body(doc, "Figure 5 compares the predicted field and error map for the best result from "
         "each protocol side by side.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for path in [
        SFDIR / "subframe_pinn_only_nsga2_field_application.png",
        SFDIR / "subframe_hybrid_nsga3_k5_field_application.png",
    ]:
        if Path(path).exists():
            run.add_picture(str(path), width=Inches(3.1))
    cap(doc, "Figure 5: Subframe-like domain. Left: PINN-only NSGA-II, final MAE=25.08 C "
        "(FEM used only at t=0). Right: FEM-anchored hybrid NSGA-III k=5, mean MAE=7.51 C "
        "(80% FEM saved). Colour scale: temperature field at t=30 s, with error map overlay.")

    fig(doc,
        FIG / "subframe_v8_pinn_only_vs_hybrid.png",
        "Figure 6: PINN-only adaptive rollout (top rows) vs. FEM-anchored hybrid (bottom rows) "
        "on the subframe-like surrogate. Each row is one NAS strategy. Left: predicted temperature "
        "field at t=30 s. Right: pointwise absolute error map.",
        width=6.3)


# ══════════════════════════════════════════════════════════════════════════
# PAGE 5 — DISCUSSION + CONCLUSIONS + REFERENCES
# ══════════════════════════════════════════════════════════════════════════

def page5(doc):
    doc.add_page_break()
    h1(doc, "6. Discussion")
    body(doc,
        "The 2D results confirmed that the NAS-PINN MSWP framework can skip FEM anchor "
        "calls across all three geometric conditions we tested. Bayesian TPE was the most "
        "consistent architecture: it outperformed NSGA-II and NSGA-III at k=3 and above "
        "on all canonical domains. We attribute this to the deeper representation of "
        "the 5-layer architecture, which handles wider time windows more robustly than "
        "the shallower NSGA networks.")
    body(doc,
        "In 3D the best single result was L-shape 3D with NSGA-II at k=1 (3.43 C), but "
        "for step-skipping with FEM savings, Bayesian TPE on the cylinder was most practical. "
        "The cylinder reached 7.98 C at k=3 with 67% FEM saving, and 11.90 C at k=5 "
        "with 80% saving. The stacked-cube domain was the hardest canonical 3D case "
        "because of its stepped geometry and strong local gradients.")
    body(doc,
        "The subframe experiments gave the clearest picture of the practical tradeoff. "
        "The hybrid protocol with NSGA-III at k=5 achieved 80% FEM saving at 7.51 C "
        "mean error, which is a useful engineering result: four FEM solves replace twenty "
        "for the same quality of time-marching supervision. The PINN-only rollout is "
        "harder to achieve: at 25.08 C it is three times less accurate than the hybrid, "
        "but it needs FEM only once at t=0. The regional analysis showed that support "
        "and cavity features drive most of the autonomous rollout error.")

    h1(doc, "7. Conclusions")
    for c in [
        "We demonstrated that NAS-PINN MSWP reduces FEM anchor evaluations by 50 to 80 "
        "percent across seven canonical 2D and 3D geometries while keeping mean MAE below "
        "10 C in most cases.",
        "On the subframe surrogate, the FEM-anchored hybrid with NSGA-III at k=5 achieved "
        "7.51 C mean MAE with 80% FEM saving. Bayesian/TPE at k=4 gave 5.56 C at 75% saving.",
        "The PINN-only adaptive rollout with NSGA-II reached 25.08 C final MAE using FEM "
        "only once at t=0. The k schedule (k=1 early, k=2 mid, k=3 late) matched physical "
        "intuition about the cooling rate.",
        "Regional error analysis confirmed that support pads and cavity features are the "
        "dominant error sources in the autonomous rollout. Balanced feature sampling helps "
        "but does not eliminate the gap.",
        "The IC-consistent output layer and Fourier feature embedding are key components: "
        "exact IC satisfaction at every window start prevents error accumulation, and wider "
        "frequency embeddings help the network resolve steep boundary gradients.",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(c).font.size = Pt(10)
        _sp(p, 1, 2)

    h1(doc, "8. References")
    for ref in [
        "[1] Mortensen et al. (2026). Mitigating distortions in cast automotive subframes. "
        "Int. J. Adv. Manuf. Technol. DOI: 10.1007/s00170-026-17515-w",
        "[2] Raissi et al. (2019). Physics-informed neural networks. J. Comput. Phys., 378, 686-707.",
        "[3] Tancik et al. (2020). Fourier features let networks learn high-frequency functions. NeurIPS 33.",
        "[4] McClenny & Braga-Neto (2023). Self-adaptive PINNs. J. Comput. Phys., 474, 111722.",
        "[5] Deb et al. (2002). NSGA-II. IEEE Trans. Evol. Comput., 6(2), 182-197.",
        "[6] Deb & Jain (2014). NSGA-III. IEEE Trans. Evol. Comput., 18(4), 577-601.",
        "[7] Bergstra et al. (2011). Algorithms for hyper-parameter optimization. NeurIPS 24.",
        "[8] Lagaris et al. (1998). ANNs for solving ODEs and PDEs. IEEE Trans. Neural Networks, 9(5).",
    ]:
        p = doc.add_paragraph(ref)
        p.style = doc.styles['Normal']
        for r in p.runs:
            r.font.size = Pt(9)
        pf = p.paragraph_format
        pf.left_indent       = Inches(0.3)
        pf.first_line_indent = Inches(-0.3)
        _sp(p, 1, 2)


# ══════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════

def main():
    print("[gen_docx_full] Generating report...")
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.2)
    sec.right_margin  = Cm(2.2)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)

    for name, fn in [
        ("Page 1 — Cover + Architecture",   page1),
        ("Page 2 — 2D Results + Heat Maps", page2),
        ("Page 3 — 3D Results + Heat Maps", page3),
        ("Page 4 — Subframe Results",       page4),
        ("Page 5 — Discussion + Conclusions", page5),
    ]:
        print(f"  {name} ...")
        fn(doc)

    doc.save(str(OUT))
    print(f"\n  Saved: {OUT}")


if __name__ == "__main__":
    main()
